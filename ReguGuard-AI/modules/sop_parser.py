import re
import io
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_file(uploaded_file):
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".pdf"):
        return _extract_pdf(uploaded_file)
    elif file_name.endswith(".docx"):
        return _extract_docx(uploaded_file)
    elif file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {file_name}")


def extract_text_and_pages_from_file(uploaded_file):
    """
    Returns (full_text, pages_text)
    - full_text: str
    - pages_text: list[str] | None   (1 item per PDF page; None for non-PDF)
    """
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".pdf"):
        return _extract_pdf_with_pages(uploaded_file)
    elif file_name.endswith(".docx"):
        return _extract_docx(uploaded_file), None
    elif file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore"), None
    else:
        raise ValueError(f"Unsupported file type: {file_name}")


def extract_text_from_bytes(filename: str, file_bytes: bytes):
    """
    Extract text from raw bytes, using the filename extension to decide parser.
    Useful for comparing already-generated documents (stored in session state).
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_bytes(file_bytes)
    elif name.endswith(".docx"):
        return _extract_docx_bytes(file_bytes)
    elif name.endswith(".txt"):
        return (file_bytes or b"").decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def extract_text_and_pages_from_bytes(filename: str, file_bytes: bytes):
    """
    Returns (full_text, pages_text) from raw bytes.
    - pages_text is only available for PDFs; otherwise None.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_with_pages_bytes(file_bytes)
    elif name.endswith(".docx"):
        return _extract_docx_bytes(file_bytes), None
    elif name.endswith(".txt"):
        return (file_bytes or b"").decode("utf-8", errors="ignore"), None
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _extract_pdf(uploaded_file):
    reader = PdfReader(io.BytesIO(uploaded_file.read()))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def _extract_pdf_with_pages(uploaded_file):
    reader = PdfReader(io.BytesIO(uploaded_file.read()))
    pages = []
    full = ""
    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        pages.append(page_text)
        if page_text:
            full += page_text + "\n"
    return full, pages


def _extract_pdf_bytes(file_bytes: bytes):
    reader = PdfReader(io.BytesIO(file_bytes or b""))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def _extract_pdf_with_pages_bytes(file_bytes: bytes):
    reader = PdfReader(io.BytesIO(file_bytes or b""))
    pages = []
    full = ""
    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        pages.append(page_text)
        if page_text:
            full += page_text + "\n"
    return full, pages


def _extract_docx(uploaded_file):
    doc = Document(io.BytesIO(uploaded_file.read()))
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    return text


def _extract_docx_bytes(file_bytes: bytes):
    doc = Document(io.BytesIO(file_bytes or b""))
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    return text


def parse_sop_into_clauses(raw_text):
    if not raw_text or not raw_text.strip():
        return []

    patterns = [
        r'(?:^|\n)\s*(?:Section|SECTION|Clause|CLAUSE|Article|ARTICLE|Policy|POLICY)\s*[\d.]+[:\s]',
        r'(?:^|\n)\s*\d+\.\d+[\s.]',
        r'(?:^|\n)\s*\d+\)\s',
        r'(?:^|\n)\s*[A-Z][A-Z\s]{5,}(?:\n)',
        r'(?:^|\n)\s*(?:Purpose|Scope|Procedure|Responsibility|References|Definitions|Objective|Compliance|Policy Statement)\s*[:\n]',
    ]

    split_points = set()
    for pattern in patterns:
        for match in re.finditer(pattern, raw_text, re.IGNORECASE):
            split_points.add(match.start())

    if not split_points:
        paragraphs = raw_text.split("\n\n")
        clauses = []
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if len(para) > 30:
                clauses.append({
                    "id": f"CLAUSE-{i+1:03d}",
                    "title": _extract_title(para),
                    "text": para,
                    "category": _categorize_clause(para),
                })
        return clauses if clauses else [{
            "id": "CLAUSE-001",
            "title": "Full Document",
            "text": raw_text.strip(),
            "category": "General",
        }]

    split_points = sorted(split_points)
    if split_points[0] > 0:
        split_points.insert(0, 0)

    clauses = []
    for i in range(len(split_points)):
        start = split_points[i]
        end = split_points[i + 1] if i + 1 < len(split_points) else len(raw_text)
        chunk = raw_text[start:end].strip()
        if len(chunk) > 20:
            clauses.append({
                "id": f"CLAUSE-{len(clauses)+1:03d}",
                "title": _extract_title(chunk),
                "text": chunk,
                "category": _categorize_clause(chunk),
            })

    return clauses


def _extract_title(text):
    first_line = text.split("\n")[0].strip()
    if len(first_line) > 100:
        first_line = first_line[:97] + "..."
    return first_line if first_line else "Untitled Clause"


CATEGORY_KEYWORDS = {
    "Data Privacy": ["data", "privacy", "personal information", "gdpr", "ccpa", "pii", "consent"],
    "Health & Safety": ["health", "safety", "osha", "hazard", "ppe", "injury", "workplace safety"],
    "Financial": ["financial", "accounting", "tax", "revenue", "audit", "fiscal", "budget"],
    "Employment": ["employee", "hiring", "termination", "leave", "compensation", "hr", "human resources"],
    "Environmental": ["environmental", "waste", "emission", "pollution", "epa", "sustainability"],
    "IT Security": ["cybersecurity", "password", "encryption", "firewall", "access control", "it security"],
    "Quality Control": ["quality", "inspection", "testing", "iso", "standard", "certification"],
    "Legal": ["legal", "contract", "liability", "compliance", "regulation", "law"],
}


def _categorize_clause(text):
    text_lower = text.lower()
    best_category = "General"
    best_score = 0
    for category, keywords in CATEGORY_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > best_score:
            best_score = score
            best_category = category
    return best_category
