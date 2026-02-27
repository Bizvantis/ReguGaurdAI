"""
Build an updated SOP document: keep full original text and ADD approved suggestions
(latest regulations / missing elements) after the relevant sections. Neat format, proper spacing.
Updates in real time as each suggestion is approved.
"""
import re
import io
from docx import Document
from docx.shared import Pt

# Label for each inserted regulatory update (clear and professional)
ADDITION_HEADER = "Regulatory update (included as per compliance review)"
ADDITION_SEPARATOR = "\n\n" + "—" * 40 + "\n\n"


def _split_into_sections(text):
    """Split document text into sections by headings (Section X, 1.1, 2.1, etc.) or double newlines."""
    if not text or not text.strip():
        return []
    text = text.strip()
    patterns = [
        r'(?:\n|^)\s*(?:Section|SECTION|Clause|CLAUSE|Article|Policy)\s*[\d.:\s]+[^\n]*',
        r'(?:\n|^)\s*\d+\.\d+[\s.]*[^\n]*',
        r'(?:\n|^)\s*(?:Purpose|Scope|Procedure|Data|Health|Employment|Financial|Environmental|Quality|IT Security|Legal)[\s:]*[^\n]*',
    ]
    split_indices = [0]
    for pattern in patterns:
        for m in re.finditer(pattern, text, re.IGNORECASE):
            if m.start() > 0 and m.start() not in split_indices:
                split_indices.append(m.start())
    split_indices = sorted(set(split_indices))
    if split_indices[-1] != len(text):
        split_indices.append(len(text))

    sections = []
    for i in range(len(split_indices) - 1):
        start, end = split_indices[i], split_indices[i + 1]
        block = text[start:end].strip()
        if len(block) > 2:
            sections.append(block)
    if not sections:
        sections = [p.strip() for p in text.split("\n\n") if p.strip()]
    return sections


def _normalize_for_match(s):
    """Normalize string for fuzzy matching."""
    if not s:
        return ""
    return re.sub(r"\s+", " ", s.lower().strip())[:80]


def _find_section_index_for_area(sections, area):
    """Find the index of the section that best matches the finding 'area'."""
    area_norm = _normalize_for_match(area)
    area_words = set(w for w in area_norm.replace("/", " ").split() if len(w) > 1)
    best_idx = 0
    best_score = 0
    for i, sec in enumerate(sections):
        sec_norm = _normalize_for_match(sec[:200])
        score = sum(1 for w in area_words if w in sec_norm)
        if score > best_score:
            best_score = score
            best_idx = i
    return best_idx


def build_updated_text_with_additions(sop_text, approved_changes):
    """
    Build updated document text: KEEP full original content and ADD each approved
    suggestion after the section it relates to. Neat, readable format with proper spacing.
    Returns: single string with original + labeled additions.
    """
    if not sop_text or not sop_text.strip():
        return sop_text
    if not approved_changes:
        return _normalize_spacing(sop_text)

    sections = _split_into_sections(sop_text)
    if not sections:
        return _normalize_spacing(sop_text)

    # Map section index -> list of (area, suggestion) to insert after that section
    insertions_by_section = {}
    for key, data in approved_changes.items():
        area = (data.get("area") or key).split("_")[0].strip()
        suggestion = (data.get("suggestion") or "").strip()
        if not suggestion:
            continue
        idx = _find_section_index_for_area(sections, area)
        insertions_by_section.setdefault(idx, []).append((area, suggestion))

    # Build result: each original section, then any additions for that section (in order)
    result_parts = []
    for i, sec in enumerate(sections):
        result_parts.append(sec.strip())
        if i in insertions_by_section:
            for area, suggestion in insertions_by_section[i]:
                result_parts.append("")
                result_parts.append(ADDITION_SEPARATOR.strip())
                result_parts.append(f"{ADDITION_HEADER} — {area}")
                result_parts.append("")
                result_parts.append(suggestion.strip())
                result_parts.append("")

    return "\n\n".join(p for p in result_parts if p is not None)


def _normalize_spacing(text):
    """Ensure consistent double newlines between paragraphs and trim excess blanks."""
    if not text:
        return text
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    return "\n\n".join(paragraphs)


def apply_approved_changes_to_text(sop_text, approved_changes):
    """Alias: build updated text with additions (original + suggestions inserted)."""
    return build_updated_text_with_additions(sop_text, approved_changes)


def _docx_to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def create_docx_from_text(text):
    """
    Create a neat DOCX from the updated text (original + additions).
    Proper spacing: space_after for readability. Addition headers are styled (bold).
    """
    doc = Document()
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    for block in blocks:
        if not block:
            continue
        is_separator = block.strip().startswith("—") or block.strip() == "—" * 40
        is_addition_header = ADDITION_HEADER in block or "Regulatory update" in block

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(10)

        if is_separator:
            p.add_run("—" * 30)
            p.paragraph_format.space_after = Pt(6)
        elif is_addition_header:
            run = p.add_run(block)
            run.bold = True
            p.paragraph_format.space_after = Pt(6)
        else:
            p.add_run(block)
    return doc


def build_updated_document(original_filename, original_bytes, sop_text, approved_changes):
    """
    Produce updated document: full original text + each approved suggestion added
    in place (after the relevant section). Neat format, proper spacing.
    Same content for all formats; DOCX/TXT/PDF (as DOCX) all get original + additions.
    Updates in real time as approvals are added (build from current approved_changes).

    Returns: (bytes, suggested_download_filename, mime_type)
    """
    if not approved_changes or not sop_text:
        return None, None, None

    timestamp = __import__("datetime").datetime.now().strftime("%Y%m%d_%H%M%S")
    base = (original_filename or "SOP").rsplit(".", 1)[0]
    out_name = f"{base}_updated_{timestamp}"

    # Single source of truth: original text + additions (labeled, spaced)
    modified_text = build_updated_text_with_additions(sop_text, approved_changes)

    name_lower = (original_filename or "").lower() if original_filename else ""

    if name_lower.endswith(".txt"):
        return modified_text.encode("utf-8"), f"{out_name}.txt", "text/plain"

    # DOCX for .docx, .pdf, or no original (sample SOP)
    doc = create_docx_from_text(modified_text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), f"{out_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
